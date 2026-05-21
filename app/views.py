from django.shortcuts import render, redirect, get_object_or_404
from .models import Project, Step
from .services.jira_service import create_jira_project
from .utils import extract_steps_from_word
import requests
import os
from requests.auth import HTTPBasicAuth
from django.conf import settings
from django.http import JsonResponse
from docx import Document
from .models import AsanaProject



def index_view(request):
    return render(request, 'projects/index.html')

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Project, Step


def test_cases_view(request):
    """
    Vue pour afficher tous les projets avec leurs étapes et statuts
    """
    # Récupérer tous les projets de l'utilisateur
    projects = Project.objects.all()
    # projects = Project.objects.filter(owner=request.user).prefetch_related('steps')
    
    # Préparer les données pour chaque projet
    projects_data = []
    for project in projects:
        # Récupérer les étapes du projet
        steps = project.steps.all()
        
        # Calculer le pourcentage de complétion
        total_steps = steps.count()
        if total_steps > 0:
            completed_steps = steps.filter(status='correct').count()
            percentage = (completed_steps / total_steps) * 100
        else:
            percentage = 0
        
        projects_data.append({
            'id': project.id,
            'name': project.name,
            'key': project.key,
            'percentage': round(percentage),
            'steps': steps
        })
    
    context = {
        'projects': projects_data
    }
    
    return render(request, 'projects/test_cases.html', context)


def update_step_status(request):
    """
    Vue pour mettre à jour le statut d'une étape via AJAX
    """
    if request.method == 'POST':
        step_id = request.POST.get('step_id')
        status = request.POST.get('status')
        
        try:
            step = Step.objects.get(id=step_id)
            # Vérifier que l'utilisateur a le droit de modifier cette étape
            if step.project.owner == request.user:
                step.status = status
                step.save()
                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False, 'error': 'Permission denied'})
        except Step.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Step not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})




def create_project_view(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        key = request.POST.get('key')
        description = request.POST.get('description')

        # Create Jira project
        create_jira_project(name, key, description)

        # Save to DB
        project = Project.objects.create(name=name, key=key, description=description)
        return redirect('upload_steps', project_id=project.id)

    return render(request, 'projects/create_project.html')







import os
import requests
from django.shortcuts import render
from requests.auth import HTTPBasicAuth

ASANA_BASE_URL = "https://app.asana.com/api/1.0"


def testrail_request(method, endpoint, payload=None):
    base_url = os.getenv("TESTRAIL_BASE_URL", "").strip().rstrip("/")
    username = os.getenv("TESTRAIL_USERNAME", "").strip()
    api_key = os.getenv("TESTRAIL_API_KEY", "").strip()

    if not base_url or not username or not api_key:
        raise ValueError(
            f"Config TestRail manquante | "
            f"TESTRAIL_BASE_URL={base_url or 'vide'} | "
            f"TESTRAIL_USERNAME={username or 'vide'} | "
            f"TESTRAIL_API_KEY={'présent' if api_key else 'vide'}"
        )

    url = f"{base_url}/index.php?/api/v2/{endpoint}"
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
    }

    print("TESTRAIL METHOD:", method)
    print("TESTRAIL URL:", url)
    print("TESTRAIL PAYLOAD:", payload)
    print("TESTRAIL USERNAME:", username)

    response = requests.request(
        method=method,
        url=url,
        headers=headers,
        auth=HTTPBasicAuth(username, api_key),
        json=payload,
        timeout=30,
    )

    print("TESTRAIL STATUS:", response.status_code)
    print("TESTRAIL RESPONSE:", response.text)

    if response.status_code >= 400:
        error = requests.HTTPError(
            f"TestRail error {response.status_code}",
            response=response
        )
        raise error

    return response.json() if response.text else {}


def create_testrail_project_and_cases(project_name, project_description, testcases=None):
    if testcases is None:
        testcases = [
            {
                "title": f"{project_name} - Vérifier la création du projet",
                "steps": "Créer un projet avec un nom valide.",
                "expected": "Le projet est créé avec succès."
            },
            {
                "title": f"{project_name} - Vérifier l'affichage du projet",
                "steps": "Accéder à la liste des projets.",
                "expected": "Le projet apparaît dans la liste."
            },
            {
                "title": f"{project_name} - Vérifier la modification du projet",
                "steps": "Modifier les informations du projet.",
                "expected": "Les modifications sont enregistrées."
            },
        ]

    project_payload = {
        "name": project_name,
        "announcement": project_description or "",
        "show_announcement": True,
        "suite_mode": 1,
    }

    testrail_project = testrail_request("POST", "add_project", project_payload)
    project_id = testrail_project["id"]

    suites_resp = testrail_request("GET", f"get_suites/{project_id}")
    suites = suites_resp if isinstance(suites_resp, list) else suites_resp.get("suites", [])

    if not suites:
        raise ValueError(f"Aucune suite trouvée pour le projet TestRail {project_id}")

    suite_id = suites[0]["id"]

    section_payload = {
        "name": "Automated import from Asana",
        "description": f"Section créée automatiquement pour le projet {project_name}.",
        "suite_id": suite_id,
    }

    testrail_section = testrail_request("POST", f"add_section/{project_id}", section_payload)
    section_id = testrail_section["id"]

    created_cases = []

    for case in testcases:
        case_payload = {
            "title": case["title"],
            "custom_steps_separated": [
                {
                    "content": case.get("steps", ""),
                    "expected": case.get("expected", "")
                }
            ]
        }

        created_case = testrail_request("POST", f"add_case/{section_id}", case_payload)
        created_cases.append(created_case)

    # return {
    #     "project": testrail_project,
    #     "suite_id": suite_id,
    #     "section": testrail_section,
    #     "cases": created_cases,
    # }
    testrail_run = create_testrail_run(
        project_id=project_id,
        suite_id=suite_id,
        run_name=f"{project_name} - Automated Run",
        description=f"Run généré automatiquement pour le projet {project_name}",
    )

    run_id = testrail_run["id"]

    created_results = add_testrail_results_for_cases(
        run_id=run_id,
        cases=created_cases,
        status_id=4,
        comment="Cas de test créés automatiquement. Statut initial: Retest.",
    )

    return {
        "project": testrail_project,
        "suite_id": suite_id,
        "section": testrail_section,
        "cases": created_cases,
        "run": testrail_run,
        "results": created_results,
    }





def create_testrail_run(project_id, suite_id, run_name=None, description="Run créé automatiquement depuis Asana"):
    run_payload = {
        "suite_id": suite_id,
        "name": run_name or "Automated Run from Asana",
        "description": description,
        "include_all": True,
    }
    return testrail_request("POST", f"add_run/{project_id}", run_payload)


def add_testrail_results_for_cases(run_id, cases, status_id=3, comment="Résultat initial créé automatiquement depuis Asana"):
    created_results = []

    for case in cases:
        case_id = case["id"]
        result_payload = {
            "status_id": status_id,
            "comment": comment,
        }
        result = testrail_request("POST", f"add_result_for_case/{run_id}/{case_id}", result_payload)
        created_results.append(result)

    return created_results








ASANA_BASE_URL = "https://app.asana.com/api/1.0" 

def asana_project_create(request):
    ctx = {
        "form_name": "",
        "form_notes": "",
        "asana_post_status": None,
        "asana_post_body": "",
        "asana_get_status": None,
        "asana_get_body": "",
        "asana_url": "",
        "testrail_project_id": None,
        "testrail_case_count": 0,
        "testrail_run_id": None,
        "testrail_result_count": 0,
        "error": "",
        "success": "",
    }

    if request.method == "POST":
        ctx["form_name"] = (request.POST.get("name") or "").strip()
        ctx["form_notes"] = (request.POST.get("notes") or "").strip()

        if not ctx["form_name"]:
            ctx["error"] = "Le nom du projet est obligatoire."
            return render(request, "asana/project_form.html", ctx)

        token = os.getenv("ASANA_TOKEN", "").strip()
        workspace_gid = os.getenv("ASANA_WORKSPACE_GID", "").strip()
        team_gid = os.getenv("ASANA_TEAM_GID", "").strip()

        if not token:
            ctx["error"] = "ASANA_TOKEN manquant dans les variables d’environnement."
            return render(request, "asana/project_form.html", ctx)

        if not workspace_gid and not team_gid:
            ctx["error"] = "ASANA_WORKSPACE_GID ou ASANA_TEAM_GID manquant. Mets au moins un des deux."
            return render(request, "asana/project_form.html", ctx)

        headers = {
            "Authorization": f"Bearer {token}",
            "Accept": "application/json",
            "Content-Type": "application/json",
        }

        payload = {
            "data": {
                "name": ctx["form_name"],
                "notes": ctx["form_notes"],
            }
        }

        if team_gid:
            post_url = f"{ASANA_BASE_URL}/teams/{team_gid}/projects"
        else:
            post_url = f"{ASANA_BASE_URL}/workspaces/{workspace_gid}/projects"

        try:
            r = requests.post(post_url, headers=headers, json=payload, timeout=25)
            ctx["asana_post_status"] = r.status_code
            ctx["asana_post_body"] = r.text or ""

            if r.status_code >= 400:
                ctx["error"] = "Asana a refusé la création. Regarde le bloc DEBUG."
                return render(request, "asana/project_form.html", ctx)

            data = r.json() if r.text else {}
            project_gid = ((data.get("data") or {}).get("gid")) if isinstance(data, dict) else None

            if not project_gid:
                ctx["error"] = "Pas de gid renvoyé par Asana."
                return render(request, "asana/project_form.html", ctx)

            r2 = requests.get(
                f"{ASANA_BASE_URL}/projects/{project_gid}",
                headers=headers,
                params={"opt_fields": "name,permalink_url,workspace.name,team.name,archived,privacy_setting"},
                timeout=25,
            )
            ctx["asana_get_status"] = r2.status_code
            ctx["asana_get_body"] = r2.text or ""

            details = r2.json() if r2.text else {}
            project_data = (details.get("data") or {}) if isinstance(details, dict) else {}
            ctx["asana_url"] = project_data.get("permalink_url") or ""

            testrail_result = create_testrail_project_and_cases(
                project_name=project_data.get("name") or ctx["form_name"],
                project_description=ctx["form_notes"],
            )

            ctx["testrail_project_id"] = testrail_result["project"]["id"]
            ctx["testrail_case_count"] = len(testrail_result["cases"])
            ctx["testrail_run_id"] = testrail_result["run"]["id"]
            ctx["testrail_result_count"] = len(testrail_result["results"])
            # ctx["success"] = (
            #     f"Projet créé sur Asana ✅ et sur TestRail ✅ "
            #     f"(Project ID TestRail: {ctx['testrail_project_id']}, "
            #     f"{ctx['testrail_case_count']} test case(s) créé(s))."
            # )
            ctx["success"] = (
                f"Projet créé sur Asana ✅ et sur TestRail ✅ "
                f"(Project ID: {ctx['testrail_project_id']}, "
                f"Run ID: {ctx['testrail_run_id']}, "
                f"{ctx['testrail_case_count']} test case(s), "
                f"{ctx['testrail_result_count']} résultat(s) initiaux créé(s))."
            )
            return render(request, "asana/project_form.html", ctx)

        except ValueError as e:
            ctx["error"] = str(e)
            return render(request, "asana/project_form.html", ctx)

        # except requests.HTTPError as e:
        #     ctx["error"] = f"Erreur TestRail: {str(e)}"
        #     return render(request, "asana/project_form.html", ctx)

        except requests.HTTPError as e:
            response = getattr(e, "response", None)
            status_code = response.status_code if response is not None else "N/A"
            response_text = response.text if response is not None else "Aucune réponse"
            request_url = response.url if response is not None else "URL inconnue"

            ctx["error"] = (
                f"Erreur TestRail | "
                f"type={type(e).__name__} | "
                f"status={status_code} | "
                f"url={request_url} | "
                f"response={response_text}"
            )
            return render(request, "asana/project_form.html", ctx)

        except requests.RequestException as e:
            ctx["error"] = f"Erreur réseau: {str(e)}"
            return render(request, "asana/project_form.html", ctx)

    return render(request, "asana/project_form.html", ctx)








# ASANA_BASE_URL = "https://app.asana.com/api/1.0" 


# def asana_project_create(request):
#     ctx = {
#         "form_name": "",
#         "form_notes": "",
#         "asana_post_status": None,
#         "asana_post_body": "",
#         "asana_get_status": None,
#         "asana_get_body": "",
#         "asana_url": "",
#         "error": "",
#         "success": "",
#     }

#     if request.method == "POST":
#         ctx["form_name"] = (request.POST.get("name") or "").strip()
#         ctx["form_notes"] = (request.POST.get("notes") or "").strip()

#         if not ctx["form_name"]:
#             ctx["error"] = "Le nom du projet est obligatoire."
#             return render(request, "asana/project_form.html", ctx)

#         token = os.getenv("ASANA_TOKEN", "").strip()
#         workspace_gid = os.getenv("ASANA_WORKSPACE_GID", "").strip()
#         team_gid = os.getenv("ASANA_TEAM_GID", "").strip()

#         if not token:
#             ctx["error"] = "ASANA_TOKEN manquant dans les variables d’environnement."
#             return render(request, "asana/project_form.html", ctx)

#         if not workspace_gid and not team_gid:
#             ctx["error"] = "ASANA_WORKSPACE_GID ou ASANA_TEAM_GID manquant. Mets au moins un des deux."
#             return render(request, "asana/project_form.html", ctx)

#         headers = {
#             "Authorization": f"Bearer {token}",
#             "Accept": "application/json",
#             "Content-Type": "application/json",
#         }

#         payload = {"data": {"name": ctx["form_name"], "notes": ctx["form_notes"]}}

#         if team_gid:
#             post_url = f"{ASANA_BASE_URL}/teams/{team_gid}/projects"
#         else:
#             post_url = f"{ASANA_BASE_URL}/workspaces/{workspace_gid}/projects"

#         try:
#             r = requests.post(post_url, headers=headers, json=payload, timeout=25)
#             ctx["asana_post_status"] = r.status_code
#             ctx["asana_post_body"] = r.text or ""

#             print("ASANA POST URL:", post_url)
#             print("ASANA POST STATUS:", r.status_code)
#             print("ASANA POST BODY:", r.text)

#             if r.status_code >= 400:
#                 ctx["error"] = "Asana a refusé la création. Regarde le bloc DEBUG (POST BODY)."
#                 return render(request, "asana/project_form.html", ctx)

#             data = r.json() if r.text else {}

#             project_gid = ((data.get("data") or {}).get("gid")) if isinstance(data, dict) else None
#             if not project_gid:
#                 ctx["error"] = "Pas de gid renvoyé par Asana. Regarde le POST BODY dans DEBUG."
#                 return render(request, "asana/project_form.html", ctx)

#             r2 = requests.get(
#                 f"{ASANA_BASE_URL}/projects/{project_gid}",
#                 headers=headers,
#                 params={"opt_fields": "name,permalink_url,workspace.name,team.name,archived,privacy_setting"},
#                 timeout=25,
#             )
#             ctx["asana_get_status"] = r2.status_code
#             ctx["asana_get_body"] = r2.text or ""

#             print("ASANA GET STATUS:", r2.status_code)
#             print("ASANA GET BODY:", r2.text)

#             details = r2.json() if r2.text else {}
#             project_data = (details.get("data") or {}) if isinstance(details, dict) else {}
#             ctx["asana_url"] = (((details.get("data") or {}).get("permalink_url")) if isinstance(details, dict) else "") or ""

#             AsanaProject.objects.update_or_create(
#                 gid=project_gid,
#                 defaults={
#                     "name": project_data.get("name") or ctx["form_name"],
#                     "permalink_url": project_data.get("permalink_url") or "",
#                     "workspace_name": ((project_data.get("workspace") or {}).get("name")) or "",
#                     "team_name": ((project_data.get("team") or {}).get("name")) or "",
#                     "archived": bool(project_data.get("archived")),
#                     "privacy_setting": project_data.get("privacy_setting") or "",
#                 },
#             )

#             ok, msg = create_unittcms_project(
#                 name=ctx["form_name"],
#                 detail=ctx["form_notes"],
#             )

#             if not ok:
#                 ctx["error"] = f"Projet créé sur Asana, mais échec de création sur UnitTCMS: {msg}"
#                 return render(request, "asana/project_form.html", ctx)

#             if ctx["asana_url"]:
#                 ctx["success"] = "Projet créé ✅ (ouvre le lien Asana ci-dessous)."
#             else:
#                 ctx["success"] = "Projet créé ✅ (mais Asana n’a pas renvoyé le lien). Regarde le GET BODY."

#             return render(request, "asana/project_form.html", ctx)

#         except requests.RequestException as e:
#             ctx["error"] = f"Erreur réseau vers Asana: {str(e)}"
#             return render(request, "asana/project_form.html", ctx)

#         except ValueError:
#             ctx["error"] = "Asana a renvoyé une réponse non-JSON. Regarde le POST BODY."
#             return render(request, "asana/project_form.html", ctx)

#     return render(request, "asana/project_form.html", ctx)



def asana_projects_list(request):
    projects = JiraProject.objects.all()  # ou [] si tu n'en as pas
    asana_projects = AsanaProject.objects.all().order_by("-id")

    context = {
        "projects": projects,
        "asana_projects": asana_projects,
    }

    return render(request, "projects/jira_projects.html", context)





def list_jira_projects(request):
    url = f"{settings.JIRA_SERVER}/rest/api/3/project"
    # auth = HTTPBasicAuth(settings.JIRA_USER, settings.JIRA_API_TOKEN)
    auth = HTTPBasicAuth(settings.ATLASSIAN_EMAIL, settings.ATLASSIAN_API_TOKEN)
    headers = {"Accept": "application/json"}

    response = requests.get(url, headers=headers, auth=auth)
    projects = []

    if response.status_code == 200:
        raw_projects = response.json()
        for project in raw_projects:
            project_url = f"{settings.JIRA_SERVER}/browse/{project['key']}"
            projects.append({
                "name": project["name"],
                "key": project["key"],
                "url": project_url
            })
    else:
        print(f"Failed to fetch projects: {response.status_code} - {response.text}")

    
    asana_projects = []

    token = os.getenv("ASANA_TOKEN", "").strip()
    workspace_gid = os.getenv("ASANA_WORKSPACE_GID", "").strip()
    team_gid = os.getenv("ASANA_TEAM_GID", "").strip()

    if token and (workspace_gid or team_gid):

        asana_headers = {
            "Authorization": f"Bearer {token}",
            "Accept": "application/json",
        }

        if team_gid:
            asana_url = f"{ASANA_BASE_URL}/teams/{team_gid}/projects"
        else:
            asana_url = f"{ASANA_BASE_URL}/workspaces/{workspace_gid}/projects"

        params = {
            "opt_fields": "gid,name,permalink_url,workspace.name",
            "limit": 100
        }

        try:
            asana_response = requests.get(asana_url, headers=asana_headers, params=params, timeout=25)

            if asana_response.status_code == 200:
                raw_asana_projects = asana_response.json().get("data", [])

                for project in raw_asana_projects:
                    asana_projects.append({
                        "id": project.get("gid"),
                        "name": project.get("name"),
                        "workspace_name": (project.get("workspace") or {}).get("name"),
                        "permalink_url": project.get("permalink_url"),
                    })
            else:
                print(f"Failed to fetch Asana projects: {asana_response.status_code} - {asana_response.text}")

        except requests.RequestException as e:
            print(f"Asana error: {str(e)}")

    else:
        print("ASANA_TOKEN ou WORKSPACE/TEAM manquant")

    return render(request, 'projects/jira_projects.html', {'projects': projects, "asana_projects": asana_projects,})







import requests
from requests.auth import HTTPBasicAuth
from django.conf import settings
from django.shortcuts import render, redirect, get_object_or_404
from .models import Project, Step
from .utils import extract_steps_from_word


def to_adf(text):
    return {
        "type": "doc",
        "version": 1,
        "content": [
            {
                "type": "paragraph",
                "content": [
                    {
                        "type": "text",
                        "text": text
                    }
                ]
            }
        ]
    }


def upload_steps_view(request, project_id):
    project = get_object_or_404(Project, id=project_id)

    if request.method == 'POST' and request.FILES.get('file'):
        word_file = request.FILES['file']
        steps = extract_steps_from_word(word_file)

        # Jira API Authentication
        headers = {
            "Accept": "application/json",
            "Content-Type": "application/json"
        }
        # auth = HTTPBasicAuth(settings.JIRA_USER, settings.JIRA_API_TOKEN)
        auth = HTTPBasicAuth(settings.ATLASSIAN_EMAIL, settings.ATLASSIAN_API_TOKEN)

        # Step 1: Dynamically get Board ID from Jira project key
        board_id = None
        boards_url = f"{settings.JIRA_SERVER}/rest/agile/1.0/board?projectKeyOrId={project.key}"
        boards_response = requests.get(boards_url, headers=headers, auth=auth)

        if boards_response.status_code == 200:
            boards = boards_response.json().get("values", [])
            if boards:
                board_id = boards[0]["id"]
                print(f"[Jira] Board ID found: {board_id}")
            else:
                print("[Jira] No boards found for project key.")
        else:
            print(f"[Jira] Failed to fetch boards: {boards_response.status_code} - {boards_response.text}")

        # Step 2: Get Active Sprint ID (if any)
        sprint_id = None
        if board_id:
            sprint_url = f"{settings.JIRA_SERVER}/rest/agile/1.0/board/{board_id}/sprint"
            sprint_response = requests.get(sprint_url, headers=headers, auth=auth)

            if sprint_response.status_code == 200:
                sprints = sprint_response.json().get("values", [])
                active_sprint = next((s for s in sprints if s["state"] == "active"), None)
                if active_sprint:
                    sprint_id = active_sprint["id"]
                    print(f"[Jira] Active sprint found: {sprint_id}")
                else:
                    print("[Jira] No active sprint found.")
            else:
                print(f"[Jira] Failed to fetch sprints: {sprint_response.status_code} - {sprint_response.text}")

        # Step 3: Create steps in Django DB and sync to Jira
        for step in steps:
            # Save to Django
            step_obj = Step.objects.create(
                project=project,
                title=step['title'],
                description=step['description']
            )

            # Create Jira Issue with ADF description
            jira_url = f"{settings.JIRA_SERVER}/rest/api/3/issue"
            payload = {
                "fields": {
                    "project": {"key": project.key},
                    "summary": step_obj.title,
                    "description": to_adf(step_obj.description),
                    "issuetype": {"name": "Task"}  # Change to "Story" if needed
                }
            }

            response = requests.post(jira_url, json=payload, headers=headers, auth=auth)

            if response.status_code == 201:
                issue = response.json()
                issue_key = issue.get("key")
                print(f"[Jira] Issue created: {issue_key}")
                print(f"[Jira] View it at: {settings.JIRA_SERVER}/browse/{issue_key}")

                # Assign issue to active sprint
                if sprint_id:
                    sprint_assignment_url = f"{settings.JIRA_SERVER}/rest/agile/1.0/sprint/{sprint_id}/issue"
                    sprint_payload = {"issues": [issue_key]}
                    sprint_response = requests.post(sprint_assignment_url, json=sprint_payload, headers=headers, auth=auth)

                    print(f"[Jira] Assigned issue {issue_key} to sprint {sprint_id}: {sprint_response.status_code}")
                    print(f"[Jira] Sprint assignment response: {sprint_response.text}")
            else:
                print(f"[Jira] Failed to create issue for step '{step['title']}': {response.status_code}")
                print(f"[Jira] Response: {response.text}")

        return redirect('project_detail', project_id=project.id)

    return render(request, 'projects/upload_steps.html', {'project': project})






def project_detail_view(request, project_id):
    project = Project.objects.get(id=project_id)
    steps = project.steps.all()
    return render(request, 'projects/project_detail.html', {'project': project, 'steps': steps})



from .tools.testlink_client import TestLinkConnector
from django.contrib import messages


# def run_test(request):
#     if request.method == "POST":
#         try:
#             url = "http://localhost/testlink/lib/api/xmlrpc/v1/xmlrpc.php"
#             dev_key = "3258edc2b243bc76b34f29d5ee00632b"
#             connector = TestLinkConnector(url, dev_key)
#             connector.report_result(status="p")
#             messages.success(request, "✅ Test exécuté avec succès !")
#         except Exception as e:
#             messages.error(request, f"❌ Erreur lors du test : {str(e)}")
#         return redirect("test-page")

#     return render(request, "projects/run_test.html")



# def run_test(request):
#     if request.method == "POST":
#         jira_key = request.POST.get("jira_key")  # Récupère la clé du projet Jira depuis le formulaire

#         try:
#             url = "http://localhost/testlink/lib/api/xmlrpc/v1/xmlrpc.php"
#             dev_key = "3258edc2b243bc76b34f29d5ee00632b"

#             # Récupérer le projet depuis la DB locale
#             project = Project.objects.filter(key=jira_key).first()
#             test_plan_name = project.steps.first().title if project and project.steps.exists() else "DemoPlan"

#             connector = TestLinkConnector(url, dev_key, jira_key, test_plan_name)  # Passe la clé Jira et le nom du plan de test au constructeur
#             connector.report_result(status="p")
#             messages.success(request, f"✅ Test lancé pour le projet {jira_key} !")
#         except Exception as e:
#             messages.error(request, f"❌ Erreur lors du test du projet {jira_key} : {str(e)}")
#         return redirect("test-page")

#     # GET request: afficher tous les projets Jira
#     url = f"{settings.JIRA_SERVER}/rest/api/3/project"
#     auth = HTTPBasicAuth(settings.JIRA_USER, settings.JIRA_API_TOKEN)
#     headers = {"Accept": "application/json"}

#     response = requests.get(url, headers=headers, auth=auth)
#     projects = []

#     if response.status_code == 200:
#         raw_projects = response.json()
#         for project in raw_projects:
#             projects.append({
#                 "name": project["name"],
#                 "key": project["key"]
#             })
#     else:
#         print(f"Failed to fetch projects: {response.status_code} - {response.text}")

#     return render(request, "projects/run_test.html", {"projects": projects})



def run_test(request):
    if request.method == "POST":
        jira_key = request.POST.get("jira_key")

        try:
            # url = "http://tl.digitalglue.in/index.php"
            # url = "http://localhost/testlink/lib/api/xmlrpc/v1/xmlrpc.php"
            # dev_key = "3258edc2b243bc76b34f29d5ee00632b"

            # url = "http://testlink.learnkap.com:8888/login.php"

            url = "http://testlink.learnkap.com:8888/lib/api/xmlrpc/v1/xmlrpc.php"
            dev_key = "441781613d80c7e335c145866c173b74"

            project = Project.objects.filter(key=jira_key).first()
            if not project:
                raise Exception("Projet Jira non trouvé")

            test_plan_name = project.steps.first().title if project.steps.exists() else "DemoPlan"

            connector = TestLinkConnector(url, dev_key, jira_key, test_plan_name)
            connector.report_result(status="p")
            messages.success(request, f"✅ Test lancé pour le projet {jira_key} !")
        except Exception as e:
            messages.error(request, f"❌ Erreur lors du test du projet {jira_key} : {str(e)}")
        return redirect("test-page")

    url = f"{settings.JIRA_SERVER}/rest/api/3/project"
    # auth = HTTPBasicAuth(settings.JIRA_USER, settings.JIRA_API_TOKEN)
    auth = HTTPBasicAuth(settings.ATLASSIAN_EMAIL, settings.ATLASSIAN_API_TOKEN)
    headers = {"Accept": "application/json"}
    response = requests.get(url, headers=headers, auth=auth)
    projects = []
    if response.status_code == 200:
        raw_projects = response.json()
        for project in raw_projects:
            projects.append({"name": project["name"], "key": project["key"]})

    # Récupère tous les projets locaux
    local_projects = {p.key: p for p in Project.objects.prefetch_related("steps").all()}

    # Enrichir les projets Jira avec leurs steps locaux
    for project in projects:
        local = local_projects.get(project["key"])
        project["steps"] = local.steps.all() if local else []

        total_steps = project["steps"].count()
        correct_steps = project["steps"].filter(status="correct").count() if total_steps else 0
        project["percentage"] = int((correct_steps / total_steps) * 100) if total_steps else 0

    return render(request, "projects/run_test.html", {"projects": projects})



def update_step_status(request):
    if request.method == "POST":
        step_id = request.POST.get("step_id")
        status = request.POST.get("status")

        try:
            step = Step.objects.get(id=step_id)
            step.status = status
            step.save()
            return JsonResponse({"success": True})
        except Step.DoesNotExist:
            return JsonResponse({"success": False, "error": "Step not found"})

    return JsonResponse({"success": False, "error": "Invalid request"})



















def extract_project_and_steps_from_word(word_file):
    """
    Extrait les infos projet + steps depuis un fichier Word (.docx).
    Format attendu dans le fichier Word :
        Name: My Project
        Key: MYP
        Description: Demo project

        Step: Install | Install Python and Django
        Step: Configure | Configure settings.py
        Step: Run | Run migrations
    """
    doc = Document(word_file)

    project_data = {"name": None, "key": None, "description": None}
    steps = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.lower().startswith("name:"):
            project_data["name"] = text.split(":", 1)[1].strip()
        elif text.lower().startswith("key:"):
            project_data["key"] = text.split(":", 1)[1].strip()
        elif text.lower().startswith("description:"):
            project_data["description"] = text.split(":", 1)[1].strip()
        elif text.lower().startswith("step:"):
            # Format attendu : Step: Title | Description
            parts = text.split(":", 1)[1].split("|")
            title = parts[0].strip()
            description = parts[1].strip() if len(parts) > 1 else ""
            steps.append({"title": title, "description": description})

    return project_data, steps


def create_jira_project(name, key, description):
    """
    Exemple de création de projet Jira (API v3).
    Ici tu peux mettre ton vrai code pour créer le projet.
    """
    url = f"{settings.JIRA_SERVER}/rest/api/3/project"
    headers = {"Accept": "application/json", "Content-Type": "application/json"}
    # auth = HTTPBasicAuth(settings.JIRA_USER, settings.JIRA_API_TOKEN)
    auth = HTTPBasicAuth(settings.ATLASSIAN_EMAIL, settings.ATLASSIAN_API_TOKEN)

    payload = {
        "key": key,
        "name": name,
        "projectTypeKey": "software",
        "projectTemplateKey": "com.pyxis.greenhopper.jira:gh-scrum-template",
        "description": description,
        "leadAccountId": settings.JIRA_ACCOUNT_ID
    }

    response = requests.post(url, headers=headers, auth=auth, json=payload)
    print("[Jira] Project creation response:", response.status_code, response.text)
    return response


def upload_project_with_steps_view(request):
    if request.method == 'POST' and request.FILES.get('file'):
        word_file = request.FILES['file']
        project_data, steps = extract_project_and_steps_from_word(word_file)

        # 1. Créer le projet Jira
        create_jira_project(
            project_data["name"],
            project_data["key"],
            project_data["description"]
        )

        # 2. Sauvegarder en DB Django
        project = Project.objects.create(
            name=project_data["name"],
            key=project_data["key"],
            description=project_data["description"]
        )

        # Auth Jira
        headers = {"Accept": "application/json", "Content-Type": "application/json"}
        # auth = HTTPBasicAuth(settings.JIRA_USER, settings.JIRA_API_TOKEN)
        auth = HTTPBasicAuth(settings.ATLASSIAN_EMAIL, settings.ATLASSIAN_API_TOKEN)

        # 3. Récupérer board ID
        board_id = None
        boards_url = f"{settings.JIRA_SERVER}/rest/agile/1.0/board?projectKeyOrId={project.key}"
        boards_response = requests.get(boards_url, headers=headers, auth=auth)
        if boards_response.status_code == 200:
            boards = boards_response.json().get("values", [])
            if boards:
                board_id = boards[0]["id"]

        # 4. Récupérer sprint actif
        sprint_id = None
        if board_id:
            sprint_url = f"{settings.JIRA_SERVER}/rest/agile/1.0/board/{board_id}/sprint"
            sprint_response = requests.get(sprint_url, headers=headers, auth=auth)
            if sprint_response.status_code == 200:
                sprints = sprint_response.json().get("values", [])
                active_sprint = next((s for s in sprints if s["state"] == "active"), None)
                if active_sprint:
                    sprint_id = active_sprint["id"]

        # 5. Créer les steps
        for step in steps:
            step_obj = Step.objects.create(
                project=project,
                title=step['title'],
                description=step['description']
            )

            jira_url = f"{settings.JIRA_SERVER}/rest/api/3/issue"
            payload = {
                "fields": {
                    "project": {"key": project.key},
                    "summary": step_obj.title,
                    "description": to_adf(step_obj.description),
                    "issuetype": {"name": "Task"}
                }
            }

            response = requests.post(jira_url, json=payload, headers=headers, auth=auth)
            if response.status_code == 201:
                issue_key = response.json().get("key")

                # assigner au sprint actif
                if sprint_id:
                    sprint_assignment_url = f"{settings.JIRA_SERVER}/rest/agile/1.0/sprint/{sprint_id}/issue"
                    sprint_payload = {"issues": [issue_key]}
                    requests.post(sprint_assignment_url, json=sprint_payload, headers=headers, auth=auth)

        return redirect('project_detail', project_id=project.id)

    return render(request, 'projects/upload_project_with_steps.html')
